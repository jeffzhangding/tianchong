from docx import Document
from shutil import copy
import re

# d = Document('黄浦投资意见书.docx')

# print('==')
#
# for g in d.paragraphs:
#     print(g.text)


class TemplateReplace(object):

    def __init__(self, file_name, settings={}):
        """"""
        self.des_name = 'test.docx'
        # copy(file_name, self.des_name)
        # self.d = Document(self.des_name)
        self.d = Document(file_name)
        self.settings = settings

    def work(self):
        self.replace_text()
        self.d.save(self.des_name)

    def replace_text(self):
        """"""
        for g in self.d.paragraphs:
            text = g.text
            key_list = re.findall('{{[\s\S]+}}', text)
            for key in key_list:
                # for name, v  in self.settings:
                #     if name
                k = re.sub('[{} ]+', '', key)
                v = self.settings.get(k, None)
                if v is not None:
                    text.replace(key, v)
            g.text = text


if __name__ == '__main__':
    s = {
        't1': 'tttttttttttttt1',
        't2': 'tttttttttttttt2',
    }
    r = TemplateReplace('黄浦投资意见书.docx', settings=s)
    r.work()





